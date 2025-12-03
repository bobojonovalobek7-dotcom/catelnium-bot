import asyncio
import logging
import json
import os
import re
import time
from typing import Union, Dict, Any, Callable, Awaitable

from aiogram import Bot, Dispatcher, F, BaseMiddleware
from aiogram.filters import Command
from aiogram.types import (
    Message, ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove, 
    FSInputFile, TelegramObject
)
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT

# --- SOZLAMALAR ---
BOT_TOKEN = "8217853009:AAH7JUesU_iZsv1R55iBaQAcJ-1WI0qBmMw"  # Tokenni qo'ying
SUPER_ADMIN_ID = 5341602920  # O'zingizning ID
DB_FILE = "baza.json"
ADMINS_FILE = "admins.json"
TIMEOUT_SECONDS = 300  # 5 daqiqa

# --- VILOYATLAR VA TUMANLAR ---
REGIONS = {
    "Toshkent shahri": ["Bektemir", "Chilonzor", "Mirobod", "Mirzo Ulug'bek", "Olmazor", "Sergeli", "Shayxontohur", "Uchtepa", "Yakkasaroy", "Yangihayot", "Yashnobod", "Yunusobod"],
    "Toshkent viloyati": ["Angren sh", "Bekobod sh", "Chirchiq sh", "Nurafshon sh", "Olmaliq sh", "Ohangaron sh", "Yangiyo'l sh", "Bekobod", "Bo'ka", "Bo'stonliq", "Chinoz", "Qibray", "Ohangaron", "Oqqo'rg'on", "Parkent", "Piskent", "Quyi Chirchiq", "O'rta Chirchiq", "Yangiyo'l", "Yuqori Chirchiq", "Zangiota", "Toshkent tumani"],
    "Andijon viloyati": ["Andijon sh", "Xonobod sh", "Andijon", "Asaka", "Baliqchi", "Bo'z", "Buloqboshi", "Izboskan", "Jalaquduq", "Marhamat", "Oltinko'l", "Paxtaobod", "Shahrixon", "Ulug'nor", "Xo'jaobod", "Qo'rg'ontepa"],
    "Buxoro viloyati": ["Buxoro sh", "Kogon sh", "Buxoro", "G'ijduvon", "Jondor", "Kogon", "Olot", "Peshku", "Qorako'l", "Qorovulbozor", "Romitan", "Shofirkon", "Vobkent"],
    "Farg'ona viloyati": ["Farg'ona sh", "Marg'ilon sh", "Qo'qon sh", "Quvasoy sh", "Beshariq", "Bog'dod", "Buvayda", "Dang'ara", "Farg'ona", "Furqat", "Oltiariq", "Qo'shtepa", "Quva", "Rishton", "So'x", "Toshloq", "Uchko'priq", "Yozyovon", "O'zbekiston"],
    "Jizzax viloyati": ["Jizzax sh", "Arnasoy", "Baxmal", "Do'stlik", "Forish", "G'allaorol", "Sharof Rashidov", "Mirzacho'l", "Paxtakor", "Yangiobod", "Zafarobod", "Zarbdor", "Zomin"],
    "Xorazm viloyati": ["Urganch sh", "Xiva sh", "Bog'ot", "Gurlan", "Qo'shko'pir", "Shovot", "Urganch", "Xiva", "Xonqa", "Hazorasp", "Yangibozor", "Yangiariq"],
    "Namangan viloyati": ["Namangan sh", "Chortoq", "Chust", "Kosonsoy", "Mingbuloq", "Namangan", "Norin", "Pop", "To'raqo'rg'on", "Uchqo'rg'on", "Uychi", "Yangiqo'rg'on"],
    "Navoiy viloyati": ["Navoiy sh", "Zarafshon sh", "Karmana", "Konimex", "Navbahor", "Nurota", "Qiziltepa", "Tomdi", "Uchquduq", "Xatirchi"],
    "Qashqadaryo viloyati": ["Qarshi sh", "Shahrisabz sh", "Chiroqchi", "Dehqonobod", "G'uzor", "Kasbi", "Kitob", "Koson", "Mirishkor", "Muborak", "Nishon", "Qamashi", "Qarshi", "Shahrisabz", "Yakkabog'"],
    "Qoraqalpog'iston Respublikasi": ["Nukus sh", "Amudaryo", "Beruniy", "Chimboy", "Ellikqal'a", "Kegeyli", "Mo'ynoq", "Nukus", "Qanliko'l", "Qo'ng'irot", "Qorao'zak", "Shumanay", "Taxtako'pir", "To'rtko'l", "Xo'jayli", "Taxiatosh"],
    "Samarqand viloyati": ["Samarqand sh", "Kattaqo'rg'on sh", "Bulung'ur", "Ishtixon", "Jomboy", "Kattaqo'rg'on", "Narpay", "Nurobod", "Oqdaryo", "Pastdarg'om", "Paxtachi", "Payariq", "Qo'shrabot", "Samarqand", "Toyloq", "Urgut"],
    "Sirdaryo viloyati": ["Guliston sh", "Yangiyer sh", "Shirin sh", "Boyovut", "Guliston", "Mirzaobod", "Oqoltin", "Sardoba", "Sayxunobod", "Sirdaryo", "Xovos"],
    "Surxondaryo viloyati": ["Termiz sh", "Angor", "Bandixon", "Boysun", "Denov", "Jarqo'rg'on", "Muzrabot", "Oltinsoy", "Qiziriq", "Qumqo'rg'on", "Sariosiyo", "Sherobod", "Sho'rchi", "Termiz", "Uzun"]
}

# --- YORDAMCHI FUNKSIYALAR ---
def load_json(filename):
    if not os.path.exists(filename):
        return []
    try:
        with open(filename, 'r', encoding='utf-8') as file:
            return json.load(file)
    except json.JSONDecodeError:
        return []

def save_json(filename, data):
    with open(filename, 'w', encoding='utf-8') as file:
        json.dump(data, file, ensure_ascii=False, indent=4)

def get_admins():
    admins = load_json(ADMINS_FILE)
    if SUPER_ADMIN_ID not in admins:
        admins.append(SUPER_ADMIN_ID)
    return admins

def add_admin_to_db(admin_id):
    admins = get_admins()
    if admin_id not in admins:
        admins.append(admin_id)
        save_json(ADMINS_FILE, admins)
        return True
    return False

def del_admin_from_db(admin_id):
    admins = get_admins()
    if admin_id in admins and admin_id != SUPER_ADMIN_ID:
        admins.remove(admin_id)
        save_json(ADMINS_FILE, admins)
        return True
    return False

def generate_word_file(data):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)

    doc.add_heading("Muassasalar bo'yicha hisobot", 0)

    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ["Viloyat", "Tuman", "Turi", "Nomi", "Xonalar", "Direktor", "Tel", "Lokatsiya", "Link"]
    for i, text in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(10)

    for item in data:
        row_cells = table.add_row().cells
        values = [
            item.get('region', ''), item.get('district', ''), item.get('type', ''),
            item.get('name', ''), str(item.get('rooms', '')), item.get('director', ''),
            item.get('phone', ''), f"{item.get('latitude')}, {item.get('longitude')}",
            item.get('map_link', '')
        ]
        for i, val in enumerate(values):
            row_cells[i].text = str(val)
            row_cells[i].paragraphs[0].font.size = Pt(9)

    filename = "hisobot.docx"
    doc.save(filename)
    return filename

def is_valid_name(text):
    if not text or len(text.split()) < 2: return False
    if any(char.isdigit() for char in text): return False
    return True

def is_valid_phone(text):
    pattern = r"^[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}$"
    return re.match(pattern, text) is not None

# --- TIMEOUT MIDDLEWARE ---
class TimeoutMiddleware(BaseMiddleware):
    async def __call__(self, handler, event, data):
        state = data.get("state")
        if state:
            state_data = await state.get_data()
            last_time = state_data.get("last_time")
            if last_time and (time.time() - last_time > TIMEOUT_SECONDS):
                if isinstance(event, Message) and event.text != "/start":
                    await state.clear()
                    await event.answer("⚠️ Vaqt tugadi. Iltimos, /start bosib qaytadan boshlang.")
                    return
            await state.update_data(last_time=time.time())
        return await handler(event, data)

# --- STATES ---
class Form(StatesGroup):
    region = State()
    district = State()
    inst_type = State()
    inst_name = State()
    room_count = State()
    director = State()
    phone = State()
    location = State()
    confirmation = State()
    editing_select = State()
    editing_value = State()

# --- BOT SETUP ---
bot = Bot(token=BOT_TOKEN)
dp = Dispatcher(storage=MemoryStorage())
dp.message.middleware(TimeoutMiddleware())

# --- NAVIGATSIYA KLAVIATURALARI ---
def get_back_cancel_kb():
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="⬅️ Orqaga"), KeyboardButton(text="❌ Bekor qilish / Bosh menyu")]
        ], resize_keyboard=True
    )

# ---------------- LOGIKA ----------------

@dp.message(Command("start"))
async def cmd_start(message: Message, state: FSMContext):
    user_id = message.from_user.id
    admins = get_admins()
    await state.clear()
    await state.update_data(last_time=time.time())

    if user_id in admins:
        kb = ReplyKeyboardMarkup(keyboard=[
            [KeyboardButton(text="📥 WORD Hisobotni olish")],
            [KeyboardButton(text="📝 Ma'lumot qo'shish (Anketa)")]
        ], resize_keyboard=True)
        await message.answer(f"Admin paneli:\n/add_admin ID\n/del_admin ID", reply_markup=kb)
    else:
        await start_survey(message, state)

@dp.message(F.text == "❌ Bekor qilish / Bosh menyu")
async def cancel_process(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("Bosh menyuga qaytildi.", reply_markup=ReplyKeyboardRemove())
    await cmd_start(message, state)

# --- ADMIN ACTIONS ---
@dp.message(Command("add_admin"))
async def add_new_admin(message: Message):
    if message.from_user.id != SUPER_ADMIN_ID: return
    try:
        new_id = int(message.text.split()[1])
        if add_admin_to_db(new_id): await message.answer(f"Admin {new_id} qo'shildi.")
        else: await message.answer("Xatolik.")
    except: await message.answer("/add_admin 12345")

@dp.message(Command("del_admin"))
async def delete_admin(message: Message):
    if message.from_user.id != SUPER_ADMIN_ID: return
    try:
        del_id = int(message.text.split()[1])
        if del_admin_from_db(del_id): await message.answer(f"Admin {del_id} o'chirildi.")
        else: await message.answer("Xatolik.")
    except: await message.answer("/del_admin 12345")

@dp.message(F.text == "📥 WORD Hisobotni olish")
async def admin_get_report(message: Message):
    if message.from_user.id not in get_admins(): return
    data = load_json(DB_FILE)
    if not data:
        await message.answer("Ma'lumot yo'q.")
        return
    await message.answer("Tayyorlanmoqda...")
    filename = generate_word_file(data)
    await message.answer_document(FSInputFile(filename), caption="Hisobot")
    save_json(DB_FILE, [])
    if os.path.exists(filename): os.remove(filename)

@dp.message(F.text == "📝 Ma'lumot qo'shish (Anketa)")
async def admin_start_btn(message: Message, state: FSMContext):
    await start_survey(message, state)

# --- ANKETA BOSQICHLARI ---

async def start_survey(message: Message, state: FSMContext):
    # 0-bosqich: Viloyat tanlash
    buttons = []
    row = []
    for name in REGIONS.keys():
        row.append(KeyboardButton(text=name))
        if len(row) == 2:
            buttons.append(row)
            row = []
    if row: buttons.append(row)
    buttons.append([KeyboardButton(text="❌ Bekor qilish / Bosh menyu")]) # Bu yerda orqaga yo'q
    
    await message.answer("Viloyatni tanlang:", reply_markup=ReplyKeyboardMarkup(keyboard=buttons, resize_keyboard=True))
    await state.set_state(Form.region)

# 1. Viloyat -> Tuman
@dp.message(Form.region)
async def process_region(message: Message, state: FSMContext):
    # Bu startdan keyingi birinchi qadam, ORQAGA bosilsa startga o'tadi
    if message.text == "⬅️ Orqaga": # Agar oldinroqdan kelsa
         await cmd_start(message, state)
         return

    if message.text not in REGIONS:
        await message.answer("Iltimos, tugmalardan tanlang.")
        return

    await state.update_data(region=message.text)
    districts = REGIONS[message.text]
    
    buttons = []
    row = []
    for dist in districts:
        row.append(KeyboardButton(text=dist))
        if len(row) == 2: buttons.append(row); row = []
    if row: buttons.append(row)
    buttons.append([KeyboardButton(text="⬅️ Orqaga"), KeyboardButton(text="❌ Bekor qilish / Bosh menyu")])

    await message.answer(f"{message.text} tanlandi. Tumanni tanlang:", reply_markup=ReplyKeyboardMarkup(keyboard=buttons, resize_keyboard=True))
    await state.set_state(Form.district)

# 2. Tuman -> Turi
@dp.message(Form.district)
async def process_district(message: Message, state: FSMContext):
    if message.text == "⬅️ Orqaga":
        await start_survey(message, state) # Viloyat tanlashga qaytish
        return

    await state.update_data(district=message.text)
    buttons = [
        [KeyboardButton(text="Maktab"), KeyboardButton(text="MTT")],
        [KeyboardButton(text="Oilaviy poliklinika")],
        [KeyboardButton(text="⬅️ Orqaga"), KeyboardButton(text="❌ Bekor qilish / Bosh menyu")]
    ]
    await message.answer("Muassasa turini tanlang:", reply_markup=ReplyKeyboardMarkup(keyboard=buttons, resize_keyboard=True))
    await state.set_state(Form.inst_type)

# 3. Turi -> Nomi
@dp.message(Form.inst_type)
async def process_type(message: Message, state: FSMContext):
    if message.text == "⬅️ Orqaga":
        # Tumanni qayta tanlash uchun viloyatni bilish kerak
        data = await state.get_data()
        msg = Message(chat=message.chat, message_id=message.message_id, date=message.date, text=data['region'], from_user=message.from_user)
        await process_region(msg, state) # Region handlerini chaqirib yuboramiz sun'iy
        return

    if message.text not in ["Maktab", "MTT", "Oilaviy poliklinika"]:
        await message.answer("Tugmani bosing.")
        return

    await state.update_data(type=message.text)
    await message.answer("Muassasa raqami yoki nomini yozing (Masalan: 14-maktab):", reply_markup=get_back_cancel_kb())
    await state.set_state(Form.inst_name)

# 4. Nomi -> Xonalar
@dp.message(Form.inst_name)
async def process_name(message: Message, state: FSMContext):
    if message.text == "⬅️ Orqaga":
        data = await state.get_data()
        # Turini qayta tanlash uchun Tuman handlerini chaqiramiz sun'iy
        msg = Message(chat=message.chat, message_id=message.message_id, date=message.date, text=data['district'], from_user=message.from_user)
        await process_district(msg, state)
        return

    await state.update_data(name=message.text)
    await message.answer("O'quv xonalar sonini kiriting (faqat raqam):", reply_markup=get_back_cancel_kb())
    await state.set_state(Form.room_count)

# 5. Xonalar -> Direktor
@dp.message(Form.room_count)
async def process_rooms(message: Message, state: FSMContext):
    if message.text == "⬅️ Orqaga":
        data = await state.get_data()
        # Nomini qayta so'rash
        await message.answer("Muassasa raqami yoki nomini yozing:", reply_markup=get_back_cancel_kb())
        await state.set_state(Form.inst_name)
        return

    if not message.text.isdigit():
        await message.answer("Faqat raqam kiriting.")
        return
    
    await state.update_data(rooms=message.text)
    await message.answer("Direktorning F.I.O.sini to'liq yozing (Masalan: Bobojonov Alobek Omon o'g'li) :", reply_markup=get_back_cancel_kb())
    await state.set_state(Form.director)

# 6. Direktor -> Telefon
@dp.message(Form.director)
async def process_director(message: Message, state: FSMContext):
    if message.text == "⬅️ Orqaga":
        await message.answer("O'quv xonalar sonini kiriting (faqat raqam):", reply_markup=get_back_cancel_kb())
        await state.set_state(Form.room_count)
        return

    if not is_valid_name(message.text):
        await message.answer("Ism familiya to'liq va raqamsiz bo'lsin.")
        return

    await state.update_data(director=message.text)
    await message.answer("Telefon raqamni kiriting (+998...):", reply_markup=get_back_cancel_kb())
    await state.set_state(Form.phone)

# 7. Telefon -> Lokatsiya
@dp.message(Form.phone)
async def process_phone(message: Message, state: FSMContext):
    if message.text == "⬅️ Orqaga":
        await message.answer("Direktorning F.I.O.sini to'liq yozing:", reply_markup=get_back_cancel_kb())
        await state.set_state(Form.director)
        return

    if not is_valid_phone(message.text):
        await message.answer("Raqam noto'g'ri.")
        return

    await state.update_data(phone=message.text)
    kb = ReplyKeyboardMarkup(keyboard=[
        [KeyboardButton(text="📍 Lokatsiyani yuborish", request_location=True)],
        [KeyboardButton(text="⬅️ Orqaga"), KeyboardButton(text="❌ Bekor qilish / Bosh menyu")]
    ], resize_keyboard=True)
    await message.answer("Pastdagi tugmani bosib lokatsiyani yuboring:", reply_markup=kb)
    await state.set_state(Form.location)

# 8. Lokatsiya -> Tasdiqlash
@dp.message(Form.location)
async def process_location(message: Message, state: FSMContext):
    if message.text == "⬅️ Orqaga":
        await message.answer("Telefon raqamni kiriting (+998...):", reply_markup=get_back_cancel_kb())
        await state.set_state(Form.phone)
        return
    
    if not message.location:
        await message.answer("Iltimos, lokatsiya tugmasini bosing.")
        return

    lat = message.location.latitude
    lon = message.location.longitude
    link = f"https://www.google.com/maps?q={lat},{lon}"
    
    await state.update_data(latitude=lat, longitude=lon, map_link=link)
    await show_confirmation(message, state)

# --- TASDIQLASH ---
async def show_confirmation(message: Message, state: FSMContext):
    data = await state.get_data()
    text = (
        f"📋 <b>Tekshirish:</b>\n\n"
        f"1. Viloyat: {data.get('region')}\n"
        f"2. Tuman: {data.get('district')}\n"
        f"3. Turi: {data.get('type')}\n"
        f"4. Nomi: {data.get('name')}\n"
        f"5. Xonalar: {data.get('rooms')}\n"
        f"6. Direktor: {data.get('director')}\n"
        f"7. Tel: {data.get('phone')}\n"
        f"8. Lokatsiya: {data.get('map_link')}\n\n"
        "To'g'rimi?"
    )
    buttons = [
        [KeyboardButton(text="✅ Tasdiqlash"), KeyboardButton(text="✏️ O'zgartirish")],
        [KeyboardButton(text="⬅️ Orqaga (Lokatsiya)"), KeyboardButton(text="❌ Bekor qilish / Bosh menyu")]
    ]
    await message.answer(text, reply_markup=ReplyKeyboardMarkup(keyboard=buttons, resize_keyboard=True), parse_mode="HTML")
    await state.set_state(Form.confirmation)

@dp.message(Form.confirmation)
async def confirm_handler(message: Message, state: FSMContext):
    if message.text == "⬅️ Orqaga (Lokatsiya)":
        # Lokatsiyaga qaytish
        kb = ReplyKeyboardMarkup(keyboard=[
            [KeyboardButton(text="📍 Lokatsiyani yuborish", request_location=True)],
            [KeyboardButton(text="⬅️ Orqaga"), KeyboardButton(text="❌ Bekor qilish / Bosh menyu")]
        ], resize_keyboard=True)
        await message.answer("Lokatsiyani qayta yuboring:", reply_markup=kb)
        await state.set_state(Form.location)
        return

    if message.text == "✅ Tasdiqlash":
        data = await state.get_data()
        if 'last_time' in data: del data['last_time']
        if 'editing_key' in data: del data['editing_key']
        
        db = load_json(DB_FILE)
        db.append(data)
        save_json(DB_FILE, db)
        await message.answer("✅ Saqlandi!", reply_markup=ReplyKeyboardRemove())
        await cmd_start(message, state)
    
    elif message.text == "✏️ O'zgartirish":
        buttons = [
            [KeyboardButton(text="Viloyat"), KeyboardButton(text="Tuman")],
            [KeyboardButton(text="Turi"), KeyboardButton(text="Nomi")],
            [KeyboardButton(text="Xonalar"), KeyboardButton(text="Direktor")],
            [KeyboardButton(text="Tel")],
            [KeyboardButton(text="⬅️ Bekor qilish")]
        ]
        await message.answer("Nimani o'zgartiramiz?", reply_markup=ReplyKeyboardMarkup(keyboard=buttons, resize_keyboard=True))
        await state.set_state(Form.editing_select)

# --- EDITING ---
@dp.message(Form.editing_select)
async def select_edit(message: Message, state: FSMContext):
    if message.text == "⬅️ Bekor qilish":
        await show_confirmation(message, state)
        return
    
    mapping = {
        "Viloyat": "region", "Tuman": "district", "Turi": "type",
        "Nomi": "name", "Xonalar": "rooms", "Direktor": "director", "Tel": "phone"
    }
    key = mapping.get(message.text)
    if key:
        await state.update_data(editing_key=key)
        await message.answer(f"Yangi {message.text}ni kiriting:", reply_markup=ReplyKeyboardRemove())
        await state.set_state(Form.editing_value)
    else:
        await message.answer("Tanlang.")

@dp.message(Form.editing_value)
async def save_edit(message: Message, state: FSMContext):
    data = await state.get_data()
    key = data.get("editing_key")
    val = message.text
    
    if key == "rooms" and not val.isdigit():
        await message.answer("Raqam yozing.")
        return
    if key == "phone" and not is_valid_phone(val):
        await message.answer("Tel xato.")
        return
        
    await state.update_data({key: val})
    await message.answer("Yangilandi.")
    await show_confirmation(message, state)

# --- START ---
async def main():
    logging.basicConfig(level=logging.INFO)
    await dp.start_polling(bot)

if __name__ == "__main__":
    try: asyncio.run(main())
    except KeyboardInterrupt: pass
